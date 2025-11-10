# src/utils/word_converter.py
import os
import shutil
import html
from typing import Optional
from fastapi import UploadFile, HTTPException
from docx import Document
import pdfkit

# ------------------------------------------------------------------
# Helper paths
# ------------------------------------------------------------------
def _temp_path(upload_dir: str, filename: str) -> str:
    os.makedirs(upload_dir, exist_ok=True)
    return os.path.join(upload_dir, f"tmp_{os.urandom(4).hex()}_{filename}")

def _output_path(output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    return os.path.join(output_dir, f"converted_{os.urandom(4).hex()}.pdf")

# ------------------------------------------------------------------
# DOCX to HTML (preserves headings, bold, tables, alignment)
# ------------------------------------------------------------------
def _docx_to_html(doc: Document) -> str:
    parts = ['<html><head><meta charset="utf-8"></head><body style="font-family:Helvetica;margin:40px;">']

    for p in doc.paragraphs:
        txt = html.escape(p.text).replace("\n", "<br>")
        if not txt.strip():
            continue
        if p.style.name.startswith("Heading 1"):
            parts.append(f"<h1>{txt}</h1>")
        elif p.style.name.startswith("Heading 2"):
            parts.append(f"<h2>{txt}</h2>")
        else:
            align = {0: "left", 1: "center", 2: "right", 3: "justify"}.get(p.alignment, "left")
            bold = "font-weight:bold;" if any(r.bold for r in p.runs) else ""
            italic = "font-style:italic;" if any(r.italic for r in p.runs) else ""
            parts.append(f'<p style="text-align:{align};{bold}{italic}margin:8px 0;">{txt}</p>')

    for tbl in doc.tables:
        parts.append('<table style="width:100%;border-collapse:collapse;margin:20px 0;" border="1">')
        for row in tbl.rows:
            parts.append("<tr>")
            for cell in row.cells:
                cell_txt = html.escape(cell.text).replace("\n", "<br>")
                parts.append(f'<td style="padding:8px;border:1px solid #ddd;">{cell_txt}</td>')
            parts.append("</tr>")
        parts.append("</table>")

    parts.append("</body></html>")
    return "".join(parts)

# ------------------------------------------------------------------
# Main conversion – returns path to PDF
# ------------------------------------------------------------------
def convert_docx_to_pdf(
    file_docx: UploadFile,
    upload_dir: str = "uploads",
    output_dir: str = "outputs",
    page_numbers: Optional[str] = None,
) -> str:
    if not file_docx.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Only .docx files are allowed")

    src = _temp_path(upload_dir, file_docx.filename)
    dst = _output_path(output_dir)

    try:
        # 1. Save uploaded file
        with open(src, "wb") as f:
            shutil.copyfileobj(file_docx.file, f)

        # 2. DOCX to HTML
        doc = Document(src)
        html_content = _docx_to_html(doc)

        # 3. HTML to PDF (explicit path – never fails on PATH issues)
        pdfkit.from_string(
            html_content,
            dst,
            configuration=pdfkit.configuration(
                wkhtmltopdf="/usr/local/bin/wkhtmltopdf"
            ),
        )
        return dst

    finally:
        if os.path.exists(src):
            try:
                os.remove(src)
            except OSError:
                pass